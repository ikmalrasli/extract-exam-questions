import requests
import io
from docx import Document
from docx.shared import Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT
import json

def add_content_to_cell(cell, content, level):
    """Helper function to add content to a cell."""
    # Return early if content is empty or not valid
    if not content or not isinstance(content, dict):
        return  # Exit the function if content is empty or not a dictionary

    if content["type"] == "text":
        # Add Malay text
        malay_text = content["text"]["malay"]
        # Directly set the text of the first paragraph
        cell.paragraphs[0].text = malay_text  # Set Malay text in the first paragraph
        
        # Add English text in italics
        english_text = content["text"]["english"]
        run = cell.paragraphs[0].add_run(f"\n{english_text}")  # Add English text in the same paragraph
        run.italic = True  # Set the run to italic
    elif content["type"] == "row":
        # Create a table within the cell for the row items
        if "items" in content and len(content["items"]) > 0:
            # Clear any existing content
            cell.text = ""
            # Create a table with 1 row and columns equal to number of items
            table = cell.add_table(rows=1, cols=len(content["items"]))
            # table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Process each item in the row
            for idx, item in enumerate(content["items"]):
                if item["type"] in ["diagram", "table"]:
                    cell = table.cell(0, idx)
                    cell.text = ""  # Clear any existing content
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    if "url" in item:
                        run = paragraph.add_run()
                        try:
                            # Download image from URL
                            response = requests.get(item["url"])
                            image_stream = io.BytesIO(response.content)
                            # Add the image from the stream with adjusted width
                            width = Inches(6.0 / len(content["items"]))
                            run.add_picture(image_stream, width=width)
                        except Exception as e:
                            paragraph.text = f"[{item['type'].upper()} {item['number']}]"
                            print(f"Failed to load image: {e}")
                    else:
                        paragraph.text = f"[{item['type'].upper()} {item['number']}]"

                    # Add caption below image
                    caption_para = cell.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if item["type"] == "diagram":
                        caption_para.text = f"Rajah {item['number']}"
                        caption_run = caption_para.add_run(f"\nDiagram {item['number']}")
                    else:  # table
                        caption_para.text = f"Jadual {item['number']}"
                        caption_run = caption_para.add_run(f"\nTable {item['number']}")
                    caption_run.italic = True
        else:
            cell.text = "[ROW ITEM]"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif content["type"] in ["diagram", "table"]:
        # Clear existing content
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if "url" in content:
            run = paragraph.add_run()
            try:
                # Download image from URL
                response = requests.get(content["url"])
                image_stream = io.BytesIO(response.content)
                # Add the image from the stream
                run.add_picture(image_stream, width=Inches(6))
            except Exception as e:
                paragraph.text = f"[{content['type'].upper()} {content['number']}]"
                print(f"Failed to load image: {e}")
        else:
            paragraph.text = f"[{content['type'].upper()} {content['number']}]"

        # Add caption below image
        caption_para = cell.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if content["type"] == "diagram":
            caption_para.text = f"Rajah {content['number']}"
            caption_run = caption_para.add_run(f"\nDiagram {content['number']}")
        else:  # table
            caption_para.text = f"Jadual {content['number']}"
            caption_run = caption_para.add_run(f"\nTable {content['number']}")
        caption_run.italic = True
    elif content['type'] == "answer_space":
        if content['format'] == "line":
            for i in range(content['lines']):
                if level == 'question':
                    if i == 0:
                        cell.paragraphs[0].text = "\n……………………………………………………………………………………………………………………………………………………………"
                    else:
                        cell.paragraphs[0].add_run("\n……………………………………………………………………………………………………………………………………………………………")
                elif level == 'sub_q':
                    if i == 0:
                        cell.paragraphs[0].text = "\n……………………………………………………………………………………………………………………………………………"
                    else:
                        cell.paragraphs[0].add_run("\n……………………………………………………………………………………………………………………………………………")
                        
        elif content['format'] == "blank-space":
            cell.text = '\n\n\n\n\n'
        elif content['format'] == "multiple-choice":
            for i in range(len(content.get("options", []))):
                malay_option = content["options"][i]["malay"]
                english_option = content["options"][i]["english"]
                if i == 0:
                    cell.paragraphs[0].text = f"   [ ] {malay_option}"
                else:
                    cell.paragraphs[0].add_run(f"\n   [ ] {malay_option}")
                
                run = cell.paragraphs[0].add_run(f"\n        {english_option}\n")
                run.italic = True

def add_marks_to_cell(cell, marks):
    """Helper function to add marks to a cell."""
    cell.paragraphs[0].text = f"[{marks} markah]"
    cell.paragraphs[0].add_run(f"\n[{marks}")
    label = cell.paragraphs[0].add_run(f" marks")
    cell.paragraphs[0].add_run(f"]")
    label.italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

def replace_newlines(data):
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, str):
                data[key] = value.replace("\\n", "\n")
            else:
                replace_newlines(value)
    elif isinstance(data, list):
        for item in data:
            replace_newlines(item)
            
def generate(data):
    replace_newlines(data)

    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)    # Set top margin
    section.bottom_margin = Inches(0.5)  # Set bottom margin
    section.left_margin = Inches(0.5)    # Set left margin
    section.right_margin = Inches(0.5)   # Set right margin
    
    # Fill table with nested structure
    for main_q in data["main_questions"]:
        print(f"Processing main question: {main_q['number']}")
        
        # Calculate total rows needed for the current main question
        total_rows = (len(main_q["content_flow"]) +  # Add main question content rows
                    sum(len(q.get("content_flow", [])) + (1 if "marks" in q else 0) + 
                        sum(len(sq.get("content_flow", [])) + (1 if "marks" in sq else 0) 
                            for sq in q.get("sub_questions", [])) 
                        for q in main_q["questions"]))  # Add question and sub-question content rows

        print(f"Total rows for main question {main_q['number']}: {total_rows}")

        # Create a new table for each main question
        table = doc.add_table(rows=total_rows, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.3)
            row.cells[1].width = Inches(0.5)
            row.cells[2].width = Inches(0.75)
            row.cells[3].width = Inches(6.5)

        current_row = 0

        # Handle main question content
        for index, content in enumerate(main_q["content_flow"]):
            print(f"Current row: {current_row}, Content type: {content['type']}")
            if content["type"] != "questions":
                if index == 0:
                    table.rows[current_row].cells[0].text = main_q["number"]
                main_q_content_cell = table.rows[current_row].cells[1]
                main_q_content_cell.merge(table.rows[current_row].cells[-1])
                add_content_to_cell(main_q_content_cell, content, 'main_q')
                current_row += 1

        # Handle questions
        for question in main_q["questions"]:
            print(f"Processing question: {question['number']}")
            if current_row >= total_rows:
                print("Error: current_row exceeds total_rows")
                break  # Prevent accessing out of range

            table.rows[current_row].cells[1].text = question["number"]
            
            
            for content in question.get("content_flow", []):
                q_content_cell = table.rows[current_row].cells[2]
                q_content_cell.merge(table.rows[current_row].cells[-1])
                add_content_to_cell(q_content_cell, content, 'question')
                current_row += 1
                    
            if "marks" in question and "sub_questions" not in question:
                question_marks_cell = table.rows[current_row].cells[2]
                question_marks_cell.merge(table.rows[current_row].cells[-1])
                add_marks_to_cell(question_marks_cell, question["marks"])
                current_row += 1
            
            # Handle sub-questions if they exist
            for sub_q in question.get("sub_questions", []):
                print(f"Processing sub-question: {sub_q['number']}")
                
                # Check if current_row is within the valid range
                if current_row >= total_rows:
                    print(f"Error: current_row {current_row} exceeds total_rows {total_rows}")
                    break  # Prevent accessing out of range

                # Accessing the sub-question number
                table.rows[current_row].cells[2].text = sub_q["number"]
                
                for content in sub_q.get("content_flow", []):
                    sub_q_content_cell = table.rows[current_row].cells[3]
                    
                    # Check if current_row is within the valid range before accessing cells
                    if current_row >= total_rows:
                        print(f"Error: current_row {current_row} exceeds total_rows {total_rows} before adding content")
                        break  # Prevent accessing out of range
                    
                    add_content_to_cell(sub_q_content_cell, content, 'sub_q')
                    current_row += 1

                # Check if current_row is within the valid range before accessing marks cell
                if current_row >= total_rows:
                    print(f"Error: current_row {current_row} exceeds total_rows {total_rows} before accessing marks cell")
                    break  # Prevent accessing out of range

                if "marks" in sub_q:
                    sub_q_marks_cell = table.rows[current_row].cells[3]
                    add_marks_to_cell(sub_q_marks_cell, sub_q['marks'])
                    current_row += 1
                
            if "marks" in question and "sub_questions" in question:
                question_marks_cell = table.rows[current_row].cells[2]
                question_marks_cell.merge(table.rows[current_row].cells[-1])
                add_marks_to_cell(question_marks_cell, question["marks"])
                current_row += 1

        # Add a page break after each table
        doc.add_page_break()

    # Save document
    doc.save("table.docx")
    
    print("Document saved!")
    
    return "table.docx"  # Return the path of the saved document
    